import os
import uuid
import boto3
from werkzeug.utils import secure_filename
import pypdf
import docx
from typing import List, Tuple, Optional
import datetime

from genai_models.models.processed_document import ProcessedDocument
from services.vector_store import vector_store_service

class DocumentService:
    """Service for processing and managing documents"""

    def __init__(self):
        """Initialize the document service"""
        # Skip MinIO connection if SKIP_MINIO env var is set
        if os.getenv("SKIP_MINIO", "false").lower() == "true":
            print("SKIP_MINIO is set. Mocking MinIO client for tests.")
            self.s3_client = self._create_mock_s3_client()
        else:
            # Initialize S3/MinIO client
            self.s3_client = boto3.client(
                's3',
                endpoint_url=os.getenv('MINIO_URL', 'http://localhost:9000'),
                aws_access_key_id=os.getenv('MINIO_ACCESS_KEY', 'minioadmin'),
                aws_secret_access_key=os.getenv('MINIO_SECRET_KEY', 'minioadmin'),
                region_name='us-east-1',
                config=boto3.session.Config(signature_version='s3v4')
            )
        
        self.bucket_name = os.getenv('MINIO_BUCKET', 'concepts')

        # Ensure bucket exists if not skipping MinIO
        if os.getenv("SKIP_MINIO", "false").lower() != "true":
            self._ensure_bucket_exists()
            
    def _create_mock_s3_client(self):
        """Create a mock S3 client for testing"""
        class MockS3Client:
            def __init__(self):
                self.mock_objects = {}
                self.meta = type('meta', (), {'session': type('session', (), {'close': lambda: None})()})
                
            def create_bucket(self, **kwargs):
                return {}
                
            def head_bucket(self, **kwargs):
                return {}
                
            def upload_fileobj(self, file_obj, bucket, key):
                if bucket not in self.mock_objects:
                    self.mock_objects[bucket] = {}
                self.mock_objects[bucket][key] = {
                    'content': 'mock_content',
                    'size': 100,
                    'last_modified': datetime.datetime.now()
                }
                return {}
                
            def list_objects_v2(self, **kwargs):
                bucket = kwargs.get('Bucket')
                prefix = kwargs.get('Prefix', '')
                
                if bucket not in self.mock_objects or not self.mock_objects[bucket]:
                    return {}
                    
                contents = []
                for key, obj in self.mock_objects[bucket].items():
                    if key.startswith(prefix):
                        contents.append({
                            'Key': key,
                            'Size': obj['size'],
                            'LastModified': obj['last_modified']
                        })
                
                if contents:
                    return {'Contents': contents}
                return {}
                
            def delete_objects(self, **kwargs):
                bucket = kwargs.get('Bucket')
                objects = kwargs.get('Delete', {}).get('Objects', [])
                
                if bucket in self.mock_objects:
                    for obj in objects:
                        key = obj.get('Key')
                        if key in self.mock_objects[bucket]:
                            del self.mock_objects[bucket][key]
                
                return {'Deleted': [{'Key': obj.get('Key')} for obj in objects]}
                
            def delete_object(self, **kwargs):
                bucket = kwargs.get('Bucket')
                key = kwargs.get('Key')
                
                if bucket in self.mock_objects and key in self.mock_objects[bucket]:
                    del self.mock_objects[bucket][key]
                
                return {}
                
            def get_paginator(self, operation_name):
                class MockPaginator:
                    def __init__(self, client, operation):
                        self.client = client
                        self.operation = operation
                        
                    def paginate(self, **kwargs):
                        # For list_objects_v2, return a single page with all objects
                        if self.operation == 'list_objects_v2':
                            result = self.client.list_objects_v2(**kwargs)
                            yield result
                
                return MockPaginator(self, operation_name)
                
        return MockS3Client()

    def _ensure_bucket_exists(self):
        """Ensure the S3/MinIO bucket exists"""
        try:
            self.s3_client.head_bucket(Bucket=self.bucket_name)
        except:
            try:
                self.s3_client.create_bucket(Bucket=self.bucket_name)
                print(f"Created bucket: {self.bucket_name}")
            except Exception as e:
                print(f"Error creating bucket: {e}")

    def process_document(self, file, concept_id: str) -> ProcessedDocument:
        """Process a document file and store it in S3 and the vector database"""
        # Generate a unique document ID
        document_id = str(uuid.uuid4())

        # Secure the filename
        filename = secure_filename(file.filename)

        # Extract text from the document
        text = self._extract_text(file)

        # Store the original file in S3/MinIO
        s3_key = f"{concept_id}/{document_id}/{filename}"
        try:
            file.seek(0)  # Reset file pointer to beginning
            self.s3_client.upload_fileobj(file, self.bucket_name, s3_key)
        except Exception as e:
            print(f"Error uploading file to S3: {e}")

        # Split text into chunks
        chunks = self._split_text(text)

        # Ensure we have at least one chunk even if text extraction failed
        if not chunks:
            chunks = [f"Failed to extract meaningful text from {filename}"]

        # Store chunks in vector database
        metadatas = [{
            "document_id": document_id,
            "concept_id": concept_id,
            "filename": filename
        } for _ in chunks]

        vector_store_service.add_texts(texts=chunks, metadatas=metadatas)

        # Create and return processed document info
        # Determine document type based on file extension
        file_extension = filename.lower().split('.')[-1] if '.' in filename else ''
        doc_type = "OTHER"  # Default type
        if file_extension == 'pdf' and 'report' in filename.lower():
            doc_type = "INDUSTRY_REPORT"
        elif file_extension in ['doc', 'docx'] and 'brand' in filename.lower():
            doc_type = "BRAND_DECK"
        elif 'debrief' in filename.lower() or 'event' in filename.lower():
            doc_type = "PAST_EVENT_DEBRIEF"
        elif 'guideline' in filename.lower() or 'guide' in filename.lower():
            doc_type = "GUIDELINES"
        
        # Build S3 location URL
        s3_location = f"s3://{self.bucket_name}/{s3_key}"
        
        return ProcessedDocument(
            id=document_id,
            filename=filename,
            type=doc_type,
            status="COMPLETED",
            s3_location=s3_location,
            uploaded_at=datetime.datetime.now(),
            processed_at=datetime.datetime.now()
        )

    def _extract_text(self, file) -> str:
        """Extract text from various document formats"""
        filename = file.filename.lower()
        file.seek(0)  # Reset file pointer to beginning

        try:
            if filename.endswith('.pdf'):
                return self._extract_text_from_pdf(file)
            elif filename.endswith('.docx'):
                return self._extract_text_from_docx(file)
            elif filename.endswith('.txt'):
                return file.read().decode('utf-8')
            else:
                return f"Unsupported file format: {filename}"
        except Exception as e:
            print(f"Error extracting text: {e}")
            return f"Error processing file: {str(e)}"

    def _extract_text_from_pdf(self, file) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            pdf = pypdf.PdfReader(file)
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
        return text

    def _extract_text_from_docx(self, file) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
        return text

    def _split_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks"""
        if not text:
            return []

        chunks = []
        for i in range(0, len(text), chunk_size - overlap):
            chunk = text[i:i + chunk_size]
            if chunk:
                chunks.append(chunk)

        return chunks

    def get_documents(self, concept_id: str, status: Optional[str] = None) -> Tuple[List[ProcessedDocument], int]:
        """Get documents for a specific concept"""
        documents = []

        try:
            # List all objects in the bucket with the concept_id prefix
            response = self.s3_client.list_objects_v2(
                Bucket=self.bucket_name,
                Prefix=f"{concept_id}/"
            )

            # If no objects found, return empty list
            if 'Contents' not in response:
                return [], 0

            # Group objects by document_id (second level in the path)
            document_map = {}
            for obj in response.get('Contents', []):
                # Parse the key: concept_id/document_id/filename
                parts = obj['Key'].split('/')
                if len(parts) >= 3:
                    doc_id = parts[1]
                    filename = parts[2]

                    if doc_id not in document_map:
                        document_map[doc_id] = {
                            'id': doc_id,
                            'filename': filename,
                            'size': obj['Size'],
                            'last_modified': obj['LastModified']
                        }

            # Get document metadata from vector store if available
            for doc_id, doc_info in document_map.items():
                # Default status if not specified
                doc_status = status if status else "COMPLETED"

                # Create the document object
                document = ProcessedDocument(
                    id=doc_id,
                    filename=doc_info['filename'],
                    type="OTHER",  # Default type
                    status=doc_status,
                    s3_location=f"s3://{self.bucket_name}/{concept_id}/{doc_id}/{doc_info['filename']}",
                    uploaded_at=doc_info['last_modified'].isoformat(),
                    processed_at=doc_info['last_modified'].isoformat()
                )

                documents.append(document)

            # Get the total count from vector store for consistency
            count = vector_store_service.get_document_count(concept_id)
            if count == 0 and documents:
                # If vector store has no count but we found documents in MinIO,
                # use the document count from MinIO
                count = len(documents)

            return documents, count

        except Exception as e:
            print(f"Error retrieving documents from MinIO: {e}")
            # Fallback to vector store count
            count = vector_store_service.get_document_count(concept_id)
            return [], count

    def delete_document_by_id(self, document_id: str) -> bool:
        """Delete a document by ID"""
        # Delete from vector store
        vector_store_success = vector_store_service.delete_by_document_id(document_id)

        # Delete from MinIO
        minio_success = False
        try:
            # First, we need to find all objects with this document_id
            # Since we don't know the concept_id, we need to list all objects and filter
            paginator = self.s3_client.get_paginator('list_objects_v2')

            # Find all objects in the bucket
            objects_to_delete = []
            for page in paginator.paginate(Bucket=self.bucket_name):
                if 'Contents' not in page:
                    continue

                for obj in page['Contents']:
                    # Check if the document_id is in the key path
                    # The format is concept_id/document_id/filename
                    parts = obj['Key'].split('/')
                    if len(parts) >= 2 and parts[1] == document_id:
                        objects_to_delete.append({'Key': obj['Key']})

            # If we found objects to delete
            if objects_to_delete:
                # Delete the objects
                self.s3_client.delete_objects(
                    Bucket=self.bucket_name,
                    Delete={'Objects': objects_to_delete}
                )
                minio_success = True
                print(f"Deleted {len(objects_to_delete)} objects from MinIO for document {document_id}")
            else:
                print(f"No objects found in MinIO for document {document_id}")
                # If no objects found, consider it a success
                minio_success = True

        except Exception as e:
            print(f"Error deleting document from MinIO: {e}")
            minio_success = False

        # Return overall success status
        return vector_store_success and minio_success

    def close(self):
        """Close the S3 client connection"""
        if hasattr(self, 's3_client') and self.s3_client:
            try:
                # Close any open connections
                session = self.s3_client.meta.session
                if session:
                    session.close()
                print("S3 client connection closed")
            except Exception as e:
                print(f"Error closing S3 client connection: {e}")

    def __del__(self):
        """Destructor to ensure the connection is closed when the object is garbage collected"""
        self.close()

# Create a singleton instance
document_service = DocumentService()

# Export the functions for use in controllers
def process_document(file, concept_id: str) -> ProcessedDocument:
    return document_service.process_document(file, concept_id)

def get_documents(concept_id: str, status: Optional[str] = None) -> Tuple[List[ProcessedDocument], int]:
    return document_service.get_documents(concept_id, status)

def delete_document_by_id(document_id: str) -> bool:
    return document_service.delete_document_by_id(document_id)
